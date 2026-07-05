#!/usr/bin/env python3
"""Generate submission .docx files from Markdown sources."""

from __future__ import annotations

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
except ImportError:
    print("Installing python-docx...", file=sys.stderr)
    import subprocess

    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent


def set_normal_style(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_heading(text, level=min(level, 3))


def add_paragraph(doc: Document, text: str, bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic


def add_code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:|]+\|?$", line.strip()))


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j in range(cols):
            cell_text = row[j] if j < len(row) else ""
            table.rows[i].cells[j].text = cell_text


def md_to_docx(md_path: Path, docx_path: Path, *, image_path: Path | None = None) -> None:
    text = md_path.read_text(encoding="utf-8")
    doc = Document()
    set_normal_style(doc)

    lines = text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        line = lines[i]

        if in_code:
            if line.strip().startswith("```"):
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue

        if line.strip().startswith("```"):
            flush_table()
            in_code = True
            i += 1
            continue

        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            if not is_table_separator(line):
                table_buf.append(parse_table_row(line))
            i += 1
            continue
        else:
            flush_table()

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if img_match:
            if image_path and image_path.exists():
                doc.add_picture(str(image_path), width=Inches(6.0))
                last_p = doc.paragraphs[-1]
                last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            level = len(heading.group(1))
            add_heading(doc, heading.group(2).strip(), level)
            i += 1
            continue

        if line.strip().startswith(">"):
            quote = line.strip().lstrip(">").strip()
            add_paragraph(doc, quote, italic=True)
            i += 1
            continue

        bullet = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        if bullet:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(strip_md_inline(bullet.group(2)))
            i += 1
            continue

        num = re.match(r"^\d+\.\s+(.+)$", line)
        if num:
            p = doc.add_paragraph(style="List Number")
            p.add_run(strip_md_inline(num.group(1)))
            i += 1
            continue

        if line.strip():
            add_paragraph(doc, strip_md_inline(line.strip()))
        i += 1

    flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"Wrote {docx_path}")


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def main() -> None:
    jobs = [
        (ROOT / "README.md", ROOT / "README.docx", None),
        (
            ROOT / "docs" / "ARCHITECTURE.md",
            ROOT / "docs" / "ARCHITECTURE.docx",
            ROOT / "docs" / "architecture-layers.png",
        ),
        (ROOT / "docs" / "WRITEUP.md", ROOT / "docs" / "WRITEUP.docx", None),
    ]
    for md, docx, img in jobs:
        if not md.exists():
            print(f"Missing {md}", file=sys.stderr)
            sys.exit(1)
        md_to_docx(md, docx, image_path=img)


if __name__ == "__main__":
    main()
